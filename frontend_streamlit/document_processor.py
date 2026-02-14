"""
Document processing utilities for RAG system

Handles:
- PDF extraction and parsing
- Text file processing
- Document chunking
- Metadata extraction
"""

import os
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process and prepare documents for RAG indexing."""
    
    @staticmethod
    def process_text(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
        """
        Process plain text and chunk it.
        
        Args:
            text: Text content
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        step = chunk_size - overlap
        
        text = text.strip()
        
        for i in range(0, len(text), step):
            chunk = text[i : i + chunk_size]
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks if chunks else [text]
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Try to import pypdf
            try:
                from pypdf import PdfReader
            except ImportError:
                # Fallback for demo
                logger.warning("pypdf not installed, returning placeholder text")
                return f"[PDF Content from {Path(file_path).name}]", {
                    "source": file_path,
                    "type": "pdf",
                    "extraction_method": "placeholder"
                }
            
            reader = PdfReader(file_path)
            text = ""
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n\n[Page {page_num + 1}]\n{page_text}"
            
            metadata = {
                "source": file_path,
                "type": "pdf",
                "pages": len(reader.pages),
                "extraction_method": "pypdf"
            }
            
            return text if text else f"[Empty PDF: {Path(file_path).name}]", metadata
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            return f"[Error extracting PDF: {str(e)}]", {
                "source": file_path,
                "type": "pdf",
                "error": str(e)
            }
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            metadata = {
                "source": file_path,
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "extraction_method": "python-docx"
            }
            
            return text if text else f"[Empty DOCX: {Path(file_path).name}]", metadata
            
        except ImportError:
            logger.warning("python-docx not installed, returning placeholder")
            return f"[DOCX Content from {Path(file_path).name}]", {
                "source": file_path,
                "type": "docx",
                "extraction_method": "placeholder"
            }
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            return f"[Error extracting DOCX: {str(e)}]", {
                "source": file_path,
                "type": "docx",
                "error": str(e)
            }
    
    @staticmethod
    def process_file(file_path: str) -> Tuple[List[str], List[Dict]]:
        """
        Process any supported file and return chunks with metadata.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (chunk_list, metadata_list)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            text, metadata = DocumentProcessor.extract_from_pdf(file_path)
        elif file_ext == '.docx':
            text, metadata = DocumentProcessor.extract_from_docx(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            metadata = {
                "source": file_path,
                "type": "txt",
                "extraction_method": "direct_read"
            }
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Chunk the text
        chunks = DocumentProcessor.process_text(text)
        
        # Create metadata for each chunk
        metadata_list = [{
            **metadata,
            "chunk_index": i,
            "total_chunks": len(chunks),
            "chunk_size": len(chunk)
        } for i, chunk in enumerate(chunks)]
        
        return chunks, metadata_list
    
    @staticmethod
    def process_directory(directory: str) -> Tuple[List[str], List[Dict]]:
        """
        Process all documents in a directory.
        
        Args:
            directory: Path to directory
            
        Returns:
            Tuple of (all_chunks, all_metadata)
        """
        all_chunks = []
        all_metadata = []
        
        supported_formats = {'.pdf', '.txt', '.docx'}
        
        for file_path in Path(directory).glob('**/*'):
            if file_path.suffix.lower() in supported_formats:
                try:
                    chunks, metadata = DocumentProcessor.process_file(str(file_path))
                    all_chunks.extend(chunks)
                    all_metadata.extend(metadata)
                    logger.info(f"Processed {file_path.name}: {len(chunks)} chunks")
                except Exception as e:
                    logger.error(f"Failed to process {file_path.name}: {str(e)}")
        
        return all_chunks, all_metadata


class TextSummarizer:
    """Summarize documents and text."""
    
    @staticmethod
    def summarize_text(text: str, max_sentences: int = 3) -> str:
        """
        Generate simple extractive summary.
        
        Args:
            text: Text to summarize
            max_sentences: Maximum sentences in summary
            
        Returns:
            Summary text
        """
        sentences = text.split('.')
        
        # Score sentences by word count and position
        scored = []
        for i, sent in enumerate(sentences):
            words = sent.strip().split()
            if words:
                # Give more weight to early sentences
                score = len(words) / (i + 1)
                scored.append((score, sent.strip()))
        
        # Get top sentences
        top_sentences = sorted(scored, key=lambda x: x[0], reverse=True)[:max_sentences]
        top_sentences = sorted(top_sentences, key=lambda x: sentences.index(x[1] + '.'))
        
        summary = '. '.join([sent for _, sent in top_sentences])
        return summary + '.'
    
    @staticmethod
    def extract_keywords(text: str, num_keywords: int = 10) -> List[str]:
        """
        Extract important keywords from text.
        
        Args:
            text: Text to analyze
            num_keywords: Number of keywords to extract
            
        Returns:
            List of keywords
        """
        import re
        from collections import Counter
        
        # Simple keyword extraction: TF-IDF-like scoring
        words = re.findall(r'\b\w+\b', text.lower())
        
        # Filter stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'being',
            'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
            'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these',
            'those', 'it', 'its'
        }
        
        filtered = [w for w in words if len(w) > 2 and w not in stop_words]
        
        # Count and get top keywords
        counter = Counter(filtered)
        keywords = [word for word, _ in counter.most_common(num_keywords)]
        
        return keywords


# Initialize logging
logging.basicConfig(level=logging.INFO)
